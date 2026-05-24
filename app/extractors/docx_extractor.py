from __future__ import annotations

from io import BytesIO

from docx import Document

from app.schemas.content import ExtractionResult, TextBlock
from app.utils.text import collapse_whitespace

COMMON_FONT_KEYWORDS = (
    "arial",
    "calibri",
    "cambria",
    "courier",
    "georgia",
    "helvetica",
    "times",
    "verdana",
    "tahoma",
    "trebuchet",
    "garamond",
    "palatino",
)


class DOCXExtractor:
    def extract(self, file_bytes: bytes) -> ExtractionResult:
        document = Document(BytesIO(file_bytes))
        blocks: list[TextBlock] = []
        order = 0
        seen_fonts: set[str] = set()

        for paragraph in document.paragraphs:
            text = collapse_whitespace(paragraph.text)
            if not text:
                continue
            for run in paragraph.runs:
                font_name = collapse_whitespace(run.font.name or "").lower()
                if font_name:
                    seen_fonts.add(font_name)
            blocks.append(TextBlock(text=text, order=order))
            order += 1

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            font_name = collapse_whitespace(run.font.name or "").lower()
                            if font_name:
                                seen_fonts.add(font_name)
                text = collapse_whitespace(
                    " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                )
                if not text:
                    continue
                blocks.append(TextBlock(text=text, order=order))
                order += 1

        has_headers_footers = False
        for section in document.sections:
            header_text = " ".join(
                collapse_whitespace(paragraph.text)
                for paragraph in section.header.paragraphs
                if collapse_whitespace(paragraph.text)
            ).strip()
            footer_text = " ".join(
                collapse_whitespace(paragraph.text)
                for paragraph in section.footer.paragraphs
                if collapse_whitespace(paragraph.text)
            ).strip()
            if header_text or footer_text:
                has_headers_footers = True
                break

        has_non_standard_fonts = any(
            not any(keyword in font_name for keyword in COMMON_FONT_KEYWORDS)
            for font_name in seen_fonts
        )

        return ExtractionResult(
            blocks=blocks,
            extraction_method="python-docx",
            warnings=[],
            ocr_candidate=False,
            has_graphics=bool(document.inline_shapes),
            has_headers_footers=has_headers_footers,
            has_non_standard_fonts=has_non_standard_fonts,
        )
